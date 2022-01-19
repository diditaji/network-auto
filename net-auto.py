from netmiko import ConnectHandler
from netmiko.ssh_exception import NetMikoTimeoutException
from netmiko.ssh_exception import NetMikoAuthenticationException
import pandas as pd
from getpass import getpass, getuser
from pprint import pprint
from docx import Document
from docx.shared import Inches
import time
import logging
import os

document = Document()
document2 = Document()
document3 = Document()
document4 = Document()
document5 = Document()

# Change username using your username
platform = 'cisco_ios'
username = 'adm.didit.aji' 

# This will return current directory
path=os.path.dirname(os.path.realpath(__file__))
os.chdir(path)


try:
    ip_add_file = open('ipaddress.txt','r')
    print ("ipaddress.txt found !")
    time.sleep(1)
except IOError:
    print ("No ipaddress.txt found ! Please check.")
    time.sleep(3)
    exit()

def showenv():
    password = getpass()
    for host in ip_add_file:
            hostdev = host.strip()
            try:
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                sh_env= device.send_command('show environment all', use_textfsm=True)
                document.add_heading('Show Environment '+hostdev)
                paragraph = document.add_paragraph(sh_env)
                document.save(hostdev+'-show-env.docx')
                print(hostdev+'-show-env.docx'+' is done !')
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')
    ip_add_file.close()
def showrun():
    password = getpass()
    for host in ip_add_file:
            hostdev = host.strip()
            try:
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                sh_run= device.send_command('show running-config', use_textfsm=True)
                document2.add_heading('Show Running Config '+hostdev, 0)
                paragraph = document2.add_paragraph(sh_run)
                document2.save(hostdev+'-show-run.docx')
                print(hostdev+'-show-run.docx'+' is done !')
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')
    ip_add_file.close()
def showver():
    password = getpass()
    for host in ip_add_file:
            hostdev = host.strip()
            try:
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                sh_ver= device.send_command('show version',use_textfsm=True)
                hostname= sh_ver[0]['hostname']
                version= sh_ver[0]['version']
                serial= sh_ver[0]['serial']
                document3.add_heading('Show Version '+hostdev, 0)
                paragraph = document3.add_paragraph('Hostname : '+hostname)
                paragraph = document3.add_paragraph('Version : '+version)
                for i in serial:
                    paragraph = document3.add_paragraph('Serial : '+ i)
                document3.save(hostdev+'-show-ver.docx')
                print(hostdev+'-show-ver.docx'+' is done !')                
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')
    ip_add_file.close()
def showcpu():
    password = getpass()
    for host in ip_add_file:
            hostdev = host.strip()
            try:
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                sh_cpu= device.send_command('show process cpu history',use_textfsm=True)
                document4.add_heading('Show CPU History '+hostdev, 0)
                paragraph = document4.add_paragraph(sh_cpu)
                document4.save(hostdev+'-show-cpu.docx')
                print(hostdev+'-show-cpu.docx'+' is done !')
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')
    ip_add_file.close()

def showall():
    password = getpass()
    for host in ip_add_file:
            hostdev = host.strip()
            try:
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                
                sh_env= device.send_command('show environment', use_textfsm=True)
                document5.add_heading('Show Environment '+hostdev, 0)
                paragraph = document5.add_paragraph(sh_env)

                sh_run= device.send_command('show running-config', use_textfsm=True)
                document5.add_heading('Show Running Config '+hostdev, 0)
                paragraph = document5.add_paragraph(sh_run)

                sh_ver= device.send_command('show version',use_textfsm=True)
                hostname= sh_ver[0]['hostname']
                version= sh_ver[0]['version']
                serial= sh_ver[0]['serial']
                document5.add_heading('Show Version '+hostdev, 0)
                paragraph = document5.add_paragraph('Hostname : '+hostname)
                paragraph = document5.add_paragraph('Version : '+version)
                for i in serial:
                    paragraph = document5.add_paragraph('Serial : '+ i)

                sh_cpu= device.send_command('show process cpu history',use_textfsm=True)
                document5.add_heading('Show CPU History '+hostdev, 0)
                paragraph = document5.add_paragraph(sh_cpu)

                document5.save(hostdev+'-show-all.docx')
                print(hostdev+'-show-all.docx'+' is done !')
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')
    ip_add_file.close()

def showallversion():
    password = getpass()

    for host in ip_add_file:
            hostdev = host.strip()
            try:
                dataver = []
                device = ConnectHandler(device_type=platform, ip=hostdev, username=username, password=password)
                dataver.extend(device.send_command('show version', use_textfsm=True))
                df = pd.DataFrame(dataver)
                df.to_excel('switch-show-version-all.xlsx', index=False)
            except ValueError:
                logging.warning('Secret  password mistake ')
            except NetMikoTimeoutException:
                logging.warning( 'the device cannot be connected, please check whether the network is communicating normally ')
            except NetMikoAuthenticationException:
                logging.warning(' login failed, user name or password error ')       
    print('switch-show-version-all.xlsx'+' is done !')
    ip_add_file.close()
    
while True:
    print("\nNetwork Automation Menu : \n")
    print("1.Show environment-docx")
    print("2.Show running-config-docx")
    print("3.Show version-docx")    
    print("4.Show cpu-docx")
    print("5.Show all-docx")
    print("6.Show all-version-xlsx")
    print("7.Exit")

    choice=int(input("\nEnter your choice: "))


    if choice==1:
        showenv()
        time.sleep(3)
        
    elif choice==2:
        showrun()
        time.sleep(3)

    elif choice==3:
        showver()
        time.sleep(3)

    elif choice==4:
        showcpu()
        time.sleep(3)

    elif choice==5:
        showall()
        time.sleep(3)

    elif choice==6:
        showallversion()
        time.sleep(3)

    elif choice==7:
        print('Bye !')
        time.sleep(3)
        break
    else:
        print("Wrong Choice !")
        time.sleep(3)